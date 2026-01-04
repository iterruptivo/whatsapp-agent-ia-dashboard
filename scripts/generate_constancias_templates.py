#!/usr/bin/env python3
"""
Generate Word templates for constancias using docx-templates syntax.
These templates use {variable} placeholders that docx-templates will replace.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_margins(cell, top=0, bottom=0, left=100, right=100):
    """Set cell margins"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for attr, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{attr}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_header_table(doc):
    """Create header with logo placeholder and company info"""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Left cell - Logo placeholder
    left_cell = table.cell(0, 0)
    left_cell.width = Inches(2)
    p = left_cell.paragraphs[0]
    p.add_run('[LOGO]')  # Placeholder - will be replaced with actual image

    # Right cell - Company info
    right_cell = table.cell(0, 1)
    right_cell.width = Inches(4)
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('{razon_social}')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run('\n')
    p.add_run('RUC: {ruc}')
    p.add_run('\n')
    p.add_run('{direccion_empresa}')

    return table

def add_title(doc, title_text):
    """Add centered title"""
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

def add_constancia_intro(doc):
    """Add SE HACE CONSTAR section"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SE HACE CONSTAR:')
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()

def add_signature_area(doc):
    """Add signature area"""
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('_' * 40)

    # Name
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run('{firma_nombre}')
    run.bold = True

    # Title
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run('{firma_cargo}')

def add_footer(doc):
    """Add footer with date and location"""
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Lima, {fecha_emision}')

def create_constancia_separacion():
    """Create template for Constancia de Separacion"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header
    create_header_table(doc)

    # Title
    add_title(doc, 'CONSTANCIA DE SEPARACION')

    # Intro
    add_constancia_intro(doc)

    # Clause 1 - Client info
    p1 = doc.add_paragraph()
    p1.add_run('1. ').bold = True
    p1.add_run('Que el(la) Sr(a). ')
    run = p1.add_run('{cliente_nombre}')
    run.bold = True
    p1.add_run(', identificado(a) con DNI N° ')
    run = p1.add_run('{cliente_dni}')
    run.bold = True
    p1.add_run(', ha realizado la separacion del siguiente bien inmueble:')

    doc.add_paragraph()

    # Clause 2 - Property info
    p2 = doc.add_paragraph()
    p2.add_run('2. ').bold = True
    p2.add_run('Local comercial ')
    run = p2.add_run('{local_codigo}')
    run.bold = True
    p2.add_run(' ({local_rubro}) con un area de ')
    run = p2.add_run('{local_area} m²')
    run.bold = True
    p2.add_run(', nivel {local_nivel}, ubicado en el proyecto ')
    run = p2.add_run('{proyecto_nombre}')
    run.bold = True
    p2.add_run('.')

    doc.add_paragraph()

    # Clause 3 - Payment info
    p3 = doc.add_paragraph()
    p3.add_run('3. ').bold = True
    p3.add_run('El monto de separacion asciende a ')
    run = p3.add_run('US$ {monto_usd}')
    run.bold = True
    p3.add_run(' (')
    run = p3.add_run('{monto_usd_letras}')
    run.italic = True
    p3.add_run(') o su equivalente en soles S/ ')
    run = p3.add_run('{monto_pen}')
    run.bold = True
    p3.add_run(' ({monto_pen_letras}) al tipo de cambio S/ {tipo_cambio}.')

    doc.add_paragraph()

    # Clause 4 - Deposits (using FOR loop for docx-templates)
    p4 = doc.add_paragraph()
    p4.add_run('4. ').bold = True
    p4.add_run('Deposito(s) realizado(s):')

    doc.add_paragraph()

    # FOR loop for deposits using docx-templates syntax: FOR...IN...END-FOR
    p_for = doc.add_paragraph()
    p_for.add_run('{FOR deposito IN depositos}')

    p_item = doc.add_paragraph()
    p_item.style = 'List Bullet'
    p_item.add_run('{$deposito.fecha} - Op. {$deposito.numero_operacion} - {$deposito.moneda} {$deposito.monto}')

    p_end = doc.add_paragraph()
    p_end.add_run('{END-FOR deposito}')

    doc.add_paragraph()

    # Clause 5 - Validity
    p5 = doc.add_paragraph()
    p5.add_run('5. ').bold = True
    p5.add_run('La presente constancia tiene vigencia de {plazo_dias} dias, hasta el ')
    run = p5.add_run('{fecha_vencimiento}')
    run.bold = True
    p5.add_run(', fecha en la cual el cliente debera completar el pago de la cuota inicial.')

    doc.add_paragraph()

    # Clause 6 - Notes
    p6 = doc.add_paragraph()
    p6.add_run('6. ').bold = True
    p6.add_run('Se emite la presente constancia a solicitud del interesado para los fines que estime conveniente.')

    # Signature
    add_signature_area(doc)

    # Footer
    add_footer(doc)

    # Save
    output_path = 'templates/constancias/constancia-separacion.docx'
    doc.save(output_path)
    print(f'Template generado: {output_path}')
    return output_path

def create_constancia_abono():
    """Create template for Constancia de Abono"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header
    create_header_table(doc)

    # Title
    add_title(doc, 'CONSTANCIA DE ABONO')

    # Intro
    add_constancia_intro(doc)

    # Clause 1 - Client info
    p1 = doc.add_paragraph()
    p1.add_run('1. ').bold = True
    p1.add_run('Que el(la) Sr(a). ')
    run = p1.add_run('{cliente_nombre}')
    run.bold = True
    p1.add_run(', identificado(a) con DNI N° ')
    run = p1.add_run('{cliente_dni}')
    run.bold = True
    p1.add_run(', ha realizado un abono correspondiente al siguiente bien inmueble:')

    doc.add_paragraph()

    # Clause 2 - Property info
    p2 = doc.add_paragraph()
    p2.add_run('2. ').bold = True
    p2.add_run('Local comercial ')
    run = p2.add_run('{local_codigo}')
    run.bold = True
    p2.add_run(' ({local_rubro}) con un area de ')
    run = p2.add_run('{local_area} m²')
    run.bold = True
    p2.add_run(', ubicado en el proyecto ')
    run = p2.add_run('{proyecto_nombre}')
    run.bold = True
    p2.add_run('.')

    doc.add_paragraph()

    # Clause 3 - Abono info
    p3 = doc.add_paragraph()
    p3.add_run('3. ').bold = True
    p3.add_run('El abono realizado corresponde a:')

    doc.add_paragraph()

    # Abono details table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    rows_data = [
        ('Monto USD:', 'US$ {monto_usd} ({monto_usd_letras})'),
        ('Fecha de deposito:', '{fecha_deposito}'),
        ('Operacion bancaria:', '{numero_operacion}'),
        ('Tipo:', 'Abono a cuenta'),
    ]

    for i, (label, value) in enumerate(rows_data):
        table.cell(i, 0).text = label
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value

    doc.add_paragraph()

    # Clause 4 - Notes
    p4 = doc.add_paragraph()
    p4.add_run('4. ').bold = True
    p4.add_run('Se emite la presente constancia a solicitud del interesado para los fines que estime conveniente.')

    # Signature
    add_signature_area(doc)

    # Footer
    add_footer(doc)

    # Save
    output_path = 'templates/constancias/constancia-abono.docx'
    doc.save(output_path)
    print(f'Template generado: {output_path}')
    return output_path

def create_constancia_cancelacion():
    """Create template for Constancia de Cancelacion"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header
    create_header_table(doc)

    # Title
    add_title(doc, 'CONSTANCIA DE CANCELACION')

    # Intro
    add_constancia_intro(doc)

    # Clause 1 - Client info
    p1 = doc.add_paragraph()
    p1.add_run('1. ').bold = True
    p1.add_run('Que el(la) Sr(a). ')
    run = p1.add_run('{cliente_nombre}')
    run.bold = True
    p1.add_run(', identificado(a) con DNI N° ')
    run = p1.add_run('{cliente_dni}')
    run.bold = True
    p1.add_run(', ha CANCELADO en su totalidad el siguiente bien inmueble:')

    doc.add_paragraph()

    # Clause 2 - Property info
    p2 = doc.add_paragraph()
    p2.add_run('2. ').bold = True
    p2.add_run('Local comercial ')
    run = p2.add_run('{local_codigo}')
    run.bold = True
    p2.add_run(' ({local_rubro}) con un area de ')
    run = p2.add_run('{local_area} m²')
    run.bold = True
    p2.add_run(', ubicado en el proyecto ')
    run = p2.add_run('{proyecto_nombre}')
    run.bold = True
    p2.add_run('.')

    doc.add_paragraph()

    # Clause 3 - Sale info
    p3 = doc.add_paragraph()
    p3.add_run('3. ').bold = True
    p3.add_run('Monto total cancelado: ')
    run = p3.add_run('US$ {monto_total_usd}')
    run.bold = True
    p3.add_run(' ({monto_total_usd_letras}).')

    doc.add_paragraph()

    # Clause 4 - Payment history
    p4 = doc.add_paragraph()
    p4.add_run('4. ').bold = True
    p4.add_run('Historial de pagos realizados:')

    doc.add_paragraph()

    # FOR loop for payment history using docx-templates syntax: FOR...IN...END-FOR
    p_for = doc.add_paragraph()
    p_for.add_run('{FOR deposito IN depositos}')

    p_item = doc.add_paragraph()
    p_item.style = 'List Bullet'
    p_item.add_run('{$deposito.fecha} - {$deposito.tipo} - US$ {$deposito.monto} - Op. {$deposito.numero_operacion}')

    p_end = doc.add_paragraph()
    p_end.add_run('{END-FOR deposito}')

    doc.add_paragraph()

    # Clause 5 - Certification
    p5 = doc.add_paragraph()
    p5.add_run('5. ').bold = True
    p5.add_run('Se certifica que el cliente ha cumplido con el pago total del bien inmueble, quedando ')
    run = p5.add_run('LIBRE DE CUALQUIER OBLIGACION DE PAGO PENDIENTE')
    run.bold = True
    p5.add_run('.')

    doc.add_paragraph()

    # Clause 6 - Notes
    p6 = doc.add_paragraph()
    p6.add_run('6. ').bold = True
    p6.add_run('Se emite la presente constancia a solicitud del interesado para los fines que estime conveniente.')

    # Signature
    add_signature_area(doc)

    # Footer
    add_footer(doc)

    # Save
    output_path = 'templates/constancias/constancia-cancelacion.docx'
    doc.save(output_path)
    print(f'Template generado: {output_path}')
    return output_path

def main():
    # Change to project root
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    os.chdir(project_root)

    # Create output directory
    os.makedirs('templates/constancias', exist_ok=True)

    print('Generando templates de constancias...')
    print('=' * 50)

    # Generate all templates
    create_constancia_separacion()
    create_constancia_abono()
    create_constancia_cancelacion()

    print('=' * 50)
    print('Templates generados exitosamente!')
    print('\nVariables disponibles:')
    print('- Comunes: empresa_nombre, empresa_ruc, empresa_direccion')
    print('- Cliente: cliente_nombre, cliente_dni')
    print('- Local: local_codigo, local_area, proyecto_nombre')
    print('- Separacion: monto_usd, monto_pen, depositos[], fecha_vencimiento')
    print('- Abono: abono_*, precio_total, total_abonado, saldo_pendiente')
    print('- Cancelacion: historial_pagos[], total_pagado_usd, total_pagado_pen')

if __name__ == '__main__':
    main()
