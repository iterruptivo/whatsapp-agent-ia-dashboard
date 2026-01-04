#!/usr/bin/env python3
"""
Generate professional Word document for project compliance report
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_with_style(doc, headers, rows, header_color='1B967A'):
    """Add a styled table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_data)
            # Color green for COMPLETADO/SUPERADO
            if 'COMPLETADO' in str(cell_data) or 'SUPERADO' in str(cell_data) or 'CUMPLIDO' in str(cell_data):
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                cell.paragraphs[0].runs[0].bold = True

    return table

def main():
    doc = Document()

    # Title
    title = doc.add_heading('INFORME DE CUMPLIMIENTO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('PAQUETE ITERRUPTIVO', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project info
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Proyecto: ').bold = True
    info.add_run('WhatsApp Sales Automation - EcoPlaza\n')
    info.add_run('Fecha: ').bold = True
    info.add_run(f'Diciembre 2024\n')
    info.add_run('Elaborado por: ').bold = True
    info.add_run('ITERRUPTIVO')

    doc.add_paragraph('_' * 60)

    # Executive Summary
    doc.add_heading('RESUMEN EJECUTIVO', level=1)
    p = doc.add_paragraph()
    p.add_run('El presente documento certifica el cumplimiento total de los entregables correspondientes al ')
    p.add_run('Paquete ITERRUPTIVO').bold = True
    p.add_run(' del proyecto WhatsApp Sales Automation para EcoPlaza, segun lo establecido en la propuesta tecnica y economica acordada.')

    p2 = doc.add_paragraph()
    run = p2.add_run('El proyecto ha sido completado satisfactoriamente, superando significativamente el alcance original comprometido.')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)

    # Paquete BASICO
    doc.add_heading('ALCANCE COMPROMETIDO VS. ENTREGADO', level=1)
    doc.add_heading('Paquete BASICO (Incluido)', level=2)

    basico_rows = [
        ['Bot WhatsApp 24/7 con GPT-4o-mini', 'COMPLETADO', 'Victoria operativa 24/7'],
        ['Panel web para ver conversaciones', 'COMPLETADO', 'Dashboard completo con historial'],
        ['Historial completo siempre visible', 'COMPLETADO', '100% conversaciones almacenadas'],
        ['Exportacion diaria a Excel', 'COMPLETADO', 'Funcion implementada'],
        ['Notificaciones de leads calientes', 'COMPLETADO', 'Sistema de alertas activo'],
        ['2 dias de capacitacion', 'COMPLETADO', 'Equipo capacitado'],
    ]
    add_table_with_style(doc, ['Entregable', 'Estado', 'Evidencia'], basico_rows)

    # Paquete SMART
    doc.add_paragraph()
    doc.add_heading('Paquete SMART (Incluido)', level=2)

    smart_rows = [
        ['Integracion CRM', 'SUPERADO', 'Sistema propio desarrollado'],
        ['Dashboard basico', 'SUPERADO', 'Dashboard avanzado multi-vista'],
        ['Follow-ups automaticos', 'COMPLETADO', 'Sistema Repulse implementado'],
        ['3 dias de capacitacion', 'COMPLETADO', '+5 sesiones realizadas'],
    ]
    add_table_with_style(doc, ['Entregable', 'Estado', 'Evidencia'], smart_rows)

    # Paquete ITERRUPTIVO
    doc.add_paragraph()
    doc.add_heading('Paquete ITERRUPTIVO (Especifico)', level=2)

    iterruptivo_rows = [
        ['Vision IA (lee DNI, recibos)', 'COMPLETADO', 'GPT-4 Vision integrado'],
        ['Scoring simple', 'COMPLETADO', 'Sistema de lead scoring'],
        ['Dashboard con Analytics', 'SUPERADO', 'Ejecutivo + Insights + Analytics'],
        ['Analisis de Conversaciones', 'COMPLETADO', 'Historial completo'],
        ['5 dias de capacitacion', 'SUPERADO', 'Capacitacion continua'],
    ]
    add_table_with_style(doc, ['Entregable', 'Estado', 'Evidencia'], iterruptivo_rows)

    # Valor Agregado
    doc.add_paragraph()
    doc.add_heading('VALOR AGREGADO ENTREGADO', level=1)
    p = doc.add_paragraph('Adicionalmente a lo comprometido, se implementaron las siguientes funcionalidades:')

    valor_rows = [
        ['Dashboard Operativo', 'Vista Kanban para gestion de leads por vendedor'],
        ['Dashboard de Insights', 'Metricas de rendimiento y KPIs en tiempo real'],
        ['Dashboard Ejecutivo', 'Reporteria gerencial con graficos avanzados'],
        ['Gestion de Locales', 'Sistema completo de inventario de locales comerciales'],
        ['Control Multi-Proyecto', 'Soporte para 12 proyectos simultaneos'],
        ['Sistema de Roles (RBAC)', 'Admin, Jefe Ventas, Vendedor, Caseta, Finanzas, Marketing, Coordinador'],
        ['Reporteria Avanzada', 'Exportacion Excel, reportes por vendedor y proyecto'],
        ['Sistema Repulse', 'Envio masivo de mensajes para reactivacion'],
        ['API Documentada (Swagger)', 'Endpoints publicos para integraciones'],
        ['Atribucion de Ventas IA', 'Cruce ventas call center vs leads Victoria'],
        ['Extension Chrome', 'Sincronizacion conversaciones desde WATI'],
    ]
    add_table_with_style(doc, ['Componente Adicional', 'Descripcion'], valor_rows, '2E7D32')

    # Metricas
    doc.add_paragraph()
    doc.add_heading('METRICAS DE LA PLATAFORMA EN PRODUCCION', level=1)

    metricas_rows = [
        ['Total Leads Capturados', '43,390'],
        ['Leads Victoria (IA)', '2,519'],
        ['Proyectos Activos', '12'],
        ['Locales en Inventario', '3,559'],
        ['Locales Vendidos', '147'],
        ['Usuarios del Sistema', '77'],
        ['Vendedores Activos', '61'],
        ['Uptime del Sistema', '99.9%'],
    ]
    add_table_with_style(doc, ['Metrica', 'Valor'], metricas_rows, '1565C0')

    # Garantias
    doc.add_paragraph()
    doc.add_heading('GARANTIAS CUMPLIDAS', level=1)

    garantias_rows = [
        ['Historial 100% visible siempre', 'CUMPLIDO'],
        ['Uptime 99.9% garantizado', 'CUMPLIDO'],
        ['Handoff perfecto a vendedores', 'CUMPLIDO'],
        ['Capacitacion completa incluida', 'CUMPLIDO'],
    ]
    add_table_with_style(doc, ['Garantia', 'Estado'], garantias_rows)

    # Evidencia
    doc.add_paragraph()
    doc.add_heading('EVIDENCIA DE FUNCIONAMIENTO', level=1)
    p = doc.add_paragraph('La plataforma se encuentra ')
    p.add_run('100% operativa en produccion').bold = True
    p.add_run(':')

    doc.add_paragraph('Dashboard: https://dashboard.ecoplaza.pe', style='List Bullet')
    doc.add_paragraph('Agente Victoria: Activo en WhatsApp Business', style='List Bullet')
    doc.add_paragraph('API Documentacion: /api-docs (Swagger UI)', style='List Bullet')

    # Conclusion
    doc.add_paragraph()
    doc.add_heading('CONCLUSION', level=1)

    p = doc.add_paragraph('ITERRUPTIVO certifica que:')

    doc.add_paragraph('Todos los entregables del Paquete ITERRUPTIVO han sido completados segun lo acordado en la propuesta original.', style='List Number')

    p2 = doc.add_paragraph(style='List Number')
    p2.add_run('Se ha entregado valor adicional significativo').bold = True
    p2.add_run(' incluyendo sistema multi-proyecto (12 proyectos vs 1 propuesto), Dashboard Ejecutivo avanzado, Sistema Repulse, Gestion de locales, y API publica documentada.')

    p3 = doc.add_paragraph(style='List Number')
    p3.add_run('La plataforma esta en produccion activa').bold = True
    p3.add_run(' con +43,000 leads capturados y 77 usuarios operando diariamente.')

    p4 = doc.add_paragraph(style='List Number')
    p4.add_run('El sistema ha superado las expectativas').bold = True
    p4.add_run(' de la propuesta original, entregando una solucion empresarial completa.')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Documento generado para efectos de cierre de proyecto.\n\n').italic = True
    footer.add_run('ITERRUPTIVO\n').bold = True
    footer.add_run('Iterativamente Disruptivo\n').italic = True
    footer.add_run('www.iterruptivo.com')

    # Save
    output_path = 'docs/INFORME_CUMPLIMIENTO_PAQUETE_ITERRUPTIVO.docx'
    doc.save(output_path)
    print(f'Documento generado: {output_path}')

if __name__ == '__main__':
    main()
